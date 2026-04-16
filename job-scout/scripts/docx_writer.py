"""
docx_writer.py
Turn generated markdown CVs and application answers into Word (.docx) files.

cv_tailor.py and application_writer.py still produce the canonical text/JSON;
this module writes a matching .docx alongside so the README promise of
Word documents is actually kept.
"""

import os
import re

from docx import Document
from docx.shared import Pt


def _is_heading_line(line):
    """Detect markdown heading (#, ##, ###) or bold-only line used as a section header."""
    stripped = line.strip()
    if stripped.startswith("#"):
        return True
    # "**Name and Contact**" style — whole line wrapped in bold
    if stripped.startswith("**") and stripped.endswith("**") and stripped.count("**") == 2:
        return True
    return False


def _heading_level(line):
    stripped = line.strip()
    if stripped.startswith("####"):
        return 4
    if stripped.startswith("###"):
        return 3
    if stripped.startswith("##"):
        return 2
    if stripped.startswith("#"):
        return 1
    return 2  # bold-only lines default to H2


def _strip_heading_markers(line):
    stripped = line.strip().lstrip("#").strip()
    # Remove surrounding ** if still present
    if stripped.startswith("**") and stripped.endswith("**"):
        stripped = stripped[2:-2].strip()
    return stripped


_BOLD_RE = re.compile(r"\*\*(.+?)\*\*")
_ITALIC_RE = re.compile(r"(?<!\*)\*(?!\*)(.+?)(?<!\*)\*(?!\*)")


def _add_inline_runs(paragraph, text):
    """Write text to a paragraph, honouring **bold** and *italic* markdown inline markers."""
    cursor = 0
    # We handle bold first, then italics inside the remaining text. Simple and good enough.
    for match in _BOLD_RE.finditer(text):
        before = text[cursor:match.start()]
        if before:
            _add_italic_aware(paragraph, before)
        run = paragraph.add_run(match.group(1))
        run.bold = True
        cursor = match.end()
    tail = text[cursor:]
    if tail:
        _add_italic_aware(paragraph, tail)


def _add_italic_aware(paragraph, text):
    cursor = 0
    for match in _ITALIC_RE.finditer(text):
        before = text[cursor:match.start()]
        if before:
            paragraph.add_run(before)
        run = paragraph.add_run(match.group(1))
        run.italic = True
        cursor = match.end()
    tail = text[cursor:]
    if tail:
        paragraph.add_run(tail)


def markdown_cv_to_docx(markdown_text, output_path):
    """Convert the generated markdown CV into a Word document."""
    doc = Document()

    # Set a sensible default font
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    for raw_line in markdown_text.splitlines():
        line = raw_line.rstrip()
        if not line.strip():
            # Preserve blank lines as empty paragraphs
            doc.add_paragraph("")
            continue

        if _is_heading_line(line):
            level = _heading_level(line)
            heading_text = _strip_heading_markers(line)
            doc.add_heading(heading_text, level=min(level, 4))
            continue

        stripped = line.strip()
        if stripped.startswith(("- ", "* ")):
            bullet = doc.add_paragraph(style="List Bullet")
            _add_inline_runs(bullet, stripped[2:])
            continue

        # Numbered list — "1. ", "2. ", ...
        if re.match(r"^\d+\.\s", stripped):
            numbered = doc.add_paragraph(style="List Number")
            _add_inline_runs(numbered, re.sub(r"^\d+\.\s+", "", stripped))
            continue

        para = doc.add_paragraph()
        _add_inline_runs(para, stripped)

    os.makedirs(os.path.dirname(os.path.abspath(output_path)) or ".", exist_ok=True)
    doc.save(output_path)
    return output_path


APPLICATION_QUESTIONS = {
    "q1": "Why are you interested in this role?",
    "q2": "Why this company?",
    "q3": "Describe your most relevant experience",
    "q4": "What's your biggest weakness or area for growth?",
    "q5": "What are your salary expectations?",
}


def application_answers_to_docx(job_label, answers, output_path):
    """Convert the q1..q5 answer dict into a Word document with question headings."""
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    doc.add_heading(f"Application answers — {job_label}", level=1)

    for key in sorted(answers.keys()):
        question = APPLICATION_QUESTIONS.get(key, key)
        doc.add_heading(question, level=2)
        doc.add_paragraph(answers[key])
        doc.add_paragraph("")

    os.makedirs(os.path.dirname(os.path.abspath(output_path)) or ".", exist_ok=True)
    doc.save(output_path)
    return output_path
